import os
import json
from docx import Document

def generate_docx(title: str, sections_json_string: str) -> str:
    print(f"Tool called: generate_docx with title '{title}'")
    
    # Ensure the output directory exists
    os.makedirs("output", exist_ok=True)
    
    # Create a new Document
    doc = Document()
    doc.add_heading(title, 0)
    
    # Parse sections and add them to the document
    try:
        if sections_json_string.strip():
            sections = json.loads(sections_json_string)
            if isinstance(sections, list):
                for section in sections:
                    heading = section.get("heading", "")
                    content = section.get("content", "")
                    if heading:
                        doc.add_heading(heading, level=1)
                    if content:
                        doc.add_paragraph(content)
            else:
                doc.add_paragraph(str(sections))
        else:
            doc.add_paragraph("No content provided.")
    except json.JSONDecodeError:
        # Fallback if Gemini failed to generate valid JSON
        print("Warning: Failed to parse sections as JSON. Dumping raw text.")
        doc.add_paragraph(sections_json_string)
        
    # Save the document
    safe_title = "".join([c if c.isalnum() else "_" for c in title])
    filepath = os.path.join("output", f"{safe_title}.docx")
    doc.save(filepath)
    
    return filepath
